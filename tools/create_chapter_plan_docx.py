from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CLINiQ_chapter1_2_current_ppt_system_plan.md"
OUT = ROOT / "CLINiQ_Chapter_1_2_Presentation_Plan.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 110)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "EAF3EF"
BORDER = "D9E2EA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for margin, value in values.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("CLINiQ Chapter 1-2 Presentation Plan")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Current PowerPoint + Current System Delivery Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Main framing",
        "This presentation should be paper-first and system-supported: Chapters 1 and 2 explain the problem, concepts, framework, and scope; the current CLINiQ prototype is used as proof that the proposed solution is feasible.",
    )


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_borders(cell, "B7D4C4")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(35, 66, 44)
    r.font.size = Pt(10)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    p.runs[0].font.size = Pt(10.5)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = hdr[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(31, 58, 95)
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_borders(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.size = Pt(9.2)
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def parse_sections(text: str) -> dict[str, str]:
    sections: dict[str, str] = {}
    current = None
    lines: list[str] = []
    for line in text.splitlines():
        if line.startswith("## "):
            if current:
                sections[current] = "\n".join(lines).strip()
            current = line[3:].strip()
            lines = []
        elif current:
            lines.append(line)
    if current:
        sections[current] = "\n".join(lines).strip()
    return sections


def parse_markdown_table(block: str) -> tuple[list[str], list[list[str]]]:
    lines = [line.strip() for line in block.splitlines() if line.strip().startswith("|")]
    if len(lines) < 3:
        return [], []
    headers = [part.strip() for part in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        rows.append([part.strip() for part in line.strip("|").split("|")])
    return headers, rows


def extract_slide_sections(text: str) -> list[dict[str, str | list[str]]]:
    pattern = re.compile(r"### Slide (\d+) - ([^\n]+)\n(.*?)(?=\n### Slide \d+ - |\n## 15-Minute Timing)", re.S)
    slides = []
    for match in pattern.finditer(text):
        body = match.group(3).strip()
        slide = {"num": match.group(1), "title": match.group(2).strip(), "raw": body}
        for key in ("Purpose", "Time", "Presenter", "Transition", "Closing"):
            m = re.search(rf"{key}:\s*(.*)", body)
            if m:
                slide[key.lower()] = m.group(1).strip().strip('"')
        points_match = re.search(r"Key points:\n(.*?)(?=\nTransition:|\nClosing:|\Z)", body, re.S)
        if points_match:
            points = []
            for line in points_match.group(1).splitlines():
                line = line.strip()
                if line.startswith("- "):
                    points.append(line[2:])
            slide["points"] = points
        demo_match = re.search(r"Demo route:\n(.*?)(?=\nTransition:|\Z)", body, re.S)
        if demo_match:
            route = []
            for line in demo_match.group(1).splitlines():
                line = line.strip()
                if re.match(r"\d+\.", line):
                    route.append(re.sub(r"^\d+\.\s*", "", line))
            slide["demo_route"] = route
        slides.append(slide)
    return slides


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)
    add_title(doc)

    sections = parse_sections(text)

    doc.add_heading("Main Strategy", level=1)
    add_callout(
        doc,
        "Core message",
        "CLINiQ is a research-based clinic management prototype designed to address PLP clinic issues in records, appointments, patient prioritization, emergency response, and clinic workflow coordination.",
    )
    doc.add_paragraph(
        "Use the current system as proof-of-concept, not as the whole center of the presentation. The paper should remain the anchor: Chapter 1 defines the problem and objectives, while Chapter 2 explains the concepts and gap that justify the proposed system."
    )

    doc.add_heading("What To Prioritize", level=1)
    doc.add_heading("Prioritize", level=2)
    add_bullets(
        doc,
        [
            "Problem, background, and objectives",
            "Scope and limitation",
            "Conceptual framework",
            "Chapter 2 concepts and gap",
            "Short proof-of-concept demo connected to the objectives",
        ],
    )
    doc.add_heading("Do Not Over-Prioritize", level=2)
    add_bullets(
        doc,
        [
            "Long module-by-module demo",
            "Every CRUD feature",
            "Deep inventory or referral walkthrough",
            "Claims that are not fully implemented",
        ],
    )

    doc.add_heading("Current PPT Assessment", level=1)
    headers, rows = parse_markdown_table(sections.get("Current PPT Assessment", ""))
    add_table(doc, headers, rows, [1.2, 1.45, 3.85])

    doc.add_heading("Recommended Final Flow", level=1)
    doc.add_paragraph("Use 12 slides by merging the current demo overview and demo screens slides.")
    slide_rows = []
    for slide in extract_slide_sections(sections.get("Recommended Final Flow", "")):
        details = []
        if slide.get("points"):
            details.append("; ".join(slide["points"]))  # type: ignore[index]
        if slide.get("demo_route"):
            details.append("Demo route: " + "; ".join(slide["demo_route"]))  # type: ignore[index]
        slide_rows.append(
            [
                f"Slide {slide['num']}: {slide['title']}",
                str(slide.get("time", "")),
                str(slide.get("presenter", "")),
                str(slide.get("purpose", "")),
                " ".join(details),
            ]
        )
    add_table(
        doc,
        ["Slide", "Time", "Presenter", "Purpose", "Key points / route"],
        slide_rows,
        [1.25, 0.75, 0.9, 1.45, 2.15],
    )

    doc.add_heading("15-Minute Timing", level=1)
    headers, rows = parse_markdown_table(sections.get("15-Minute Timing", ""))
    add_table(doc, headers, rows, [4.9, 1.6])

    doc.add_heading("Speaker Distribution", level=1)
    doc.add_heading("Presenter 1", level=2)
    add_bullets(doc, ["Title", "Introduction", "Background", "Statement of the Problem", "Scope and Limitation"])
    doc.add_heading("Presenter 2", level=2)
    add_bullets(doc, ["Conceptual Framework", "Chapter 2 Concepts and Gap", "Core Features", "Risk Classification", "QR/NFC Passport and Nurse Alerting"])
    doc.add_heading("Presenter 3", level=2)
    add_bullets(doc, ["Short system demo", "Current prototype limitations", "Conclusion"])

    doc.add_heading("Demo Plan For Chapter 1-2 Defense", level=1)
    doc.add_paragraph("Keep the demo around 3 to 4 minutes. Show only the system parts that directly support Chapter 1 and Chapter 2.")
    demo_sections = re.split(r"\n### Demo Step \d+ - ", sections.get("Demo Plan For Chapter 1-2 Defense", ""))
    for chunk in demo_sections[1:]:
        title, _, body = chunk.partition("\n")
        doc.add_heading(title.strip(), level=2)
        show_match = re.search(r"Show:\n(.*?)(?=\nSay:)", body, re.S)
        if show_match:
            add_bullets(doc, [line.strip()[2:] for line in show_match.group(1).splitlines() if line.strip().startswith("- ")])
        say_match = re.search(r'Say:\n"([^"]+)"', body, re.S)
        if say_match:
            add_callout(doc, "Say", say_match.group(1).strip())

    doc.add_heading("Claims To Say Carefully", level=1)
    claims = sections.get("Claims To Say Carefully", "")
    use_match = re.search(r"Use:\n(.*?)(?=\n\nAvoid:)", claims, re.S)
    avoid_match = re.search(r"Avoid:\n(.*)", claims, re.S)
    doc.add_heading("Use", level=2)
    if use_match:
        add_bullets(doc, [line.strip()[2:].strip('"') for line in use_match.group(1).splitlines() if line.strip().startswith("- ")])
    doc.add_heading("Avoid", level=2)
    if avoid_match:
        add_bullets(doc, [line.strip()[2:].strip('"') for line in avoid_match.group(1).splitlines() if line.strip().startswith("- ")])

    doc.add_heading("Best Final Framing", level=1)
    quotes = re.findall(r'> "([^"]+)"', sections.get("Best Final Framing", ""))
    for label, quote in zip(["Near the beginning", "Before the demo", "Conclusion"], quotes):
        add_callout(doc, label, quote)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("CLINiQ Chapter 1-2 Presentation Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
