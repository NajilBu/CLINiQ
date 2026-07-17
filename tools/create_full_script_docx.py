from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CLINiQ_chapter1_2_full_presentation_script.md"
OUT = ROOT / "CLINiQ_Chapter_1_2_Full_Presentation_Script.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 110)
CALLOUT_FILL = "EAF3EF"
CALLOUT_BORDER = "B7D4C4"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=180, bottom=120, end=180) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("CLINiQ Chapter 1-2 Full Presentation Script")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Speaker script for the current PowerPoint and current CLINiQ prototype")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED


def add_callout(doc: Document, text: str, label: str = "Note") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.cell(0, 0).width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_borders(cell, CALLOUT_BORDER)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(35, 66, 44)
    r.font.size = Pt(10)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if text.endswith(":") and len(text) < 35:
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = DARK_BLUE
    else:
        p.add_run(text)


def convert_markdown_to_docx() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)
    add_title(doc)

    skip_first_title = True
    in_quote = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            in_quote = False
            continue
        if stripped == "---":
            doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            if skip_first_title:
                skip_first_title = False
                continue
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue

        if stripped.startswith("> "):
            quote = stripped[2:].strip().strip('"')
            add_callout(doc, quote, "Suggested line")
            in_quote = True
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(stripped[2:].strip())
            continue

        if stripped[:2].isdigit() and stripped[2:4] == ". ":
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(stripped[4:].strip())
            continue

        if stripped.startswith("Question:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(stripped)
            r.bold = True
            r.font.color.rgb = DARK_BLUE
            continue

        if stripped.startswith("Answer:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(stripped)
            r.bold = True
            r.font.color.rgb = RGBColor(35, 66, 44)
            continue

        if stripped.startswith("Transition:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run("Transition:")
            r.bold = True
            r.font.color.rgb = MUTED
            continue

        if not in_quote:
            add_body_paragraph(doc, stripped)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("CLINiQ Chapter 1-2 Presentation Script")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    convert_markdown_to_docx()
